"""
碳排放計算工具 v3.2 - Web版本
完整功能版本

作者: 楊勢賢
日期: 2024-12-27
"""

import streamlit as st
import pandas as pd
import numpy as np
from geopy.geocoders import Nominatim
from geopy.distance import geodesic
import requests
from docx import Document
from datetime import datetime
import matplotlib.pyplot as plt
import json
import time
from pathlib import Path
from io import BytesIO

# ========【頁面設定】========
st.set_page_config(
    page_title="碳排放計算工具 v3.2",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========【matplotlib 中文字型設定】========
plt.rcParams["font.sans-serif"] = ["Microsoft JhengHei", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False

# ========【常數定義】========
GWP_CH4 = 28
GWP_N2O = 265

FIXED_CO2_FACTORS = {
    "捷運": 0.04, "公車": 0.04, "大客車": 0.04, "高鐵": 0.04, "火車": 0.06,
    "摩托車": 0.046, "電動機車": 0.025, "電動車": 0.078, "飛機": 2.1981, "船": 2.606
}

RECOMMENDED_FUEL_TYPES = {
    "汽車": "車用汽油", "船": "柴油", "飛機": "航空汽油",
    "摩托車": "車用汽油", "公車": "柴油", "大客車": "柴油"
}

# ========【持續學習系統】========
class ContinuousLearningSystem:
    """持續學習系統 - 讓AI越用越聰明"""
    
    def __init__(self, storage_dir="./carbon_learning_data"):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
        # 儲存文件路徑
        self.conversation_file = self.storage_dir / "conversations.json"
        self.company_profile_file = self.storage_dir / "company_profile.json"
        self.qa_database_file = self.storage_dir / "qa_database.json"
        self.usage_stats_file = self.storage_dir / "usage_stats.json"
        
        # 載入歷史數據
        self.conversations = self._load_json(self.conversation_file, [])
        self.company_profile = self._load_json(self.company_profile_file, {})
        self.qa_database = self._load_json(self.qa_database_file, [])
        self.usage_stats = self._load_json(self.usage_stats_file, {
            "total_runs": 0,
            "total_conversations": 0,
            "first_use": None,
            "last_use": None
        })
        
        # 更新使用統計
        self.usage_stats["total_runs"] += 1
        self.usage_stats["last_use"] = datetime.now().isoformat()
        if self.usage_stats["first_use"] is None:
            self.usage_stats["first_use"] = datetime.now().isoformat()
        
        self._save_json(self.usage_stats_file, self.usage_stats)
    
    def _load_json(self, filepath, default):
        """載入JSON文件"""
        try:
            if filepath.exists():
                with open(filepath, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except:
            pass
        return default
    
    def _save_json(self, filepath, data):
        """儲存JSON文件"""
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"儲存失敗: {e}")
    
    def analyze_data(self, detailed_results):
        """分析Excel數據並更新公司檔案"""
        if detailed_results.empty:
            return
        
        analysis = {
            "last_updated": datetime.now().isoformat(),
            "total_employees": len(detailed_results),
            "units": {},
            "transport_modes": {},
            "total_emissions": float(detailed_results['總排放量(kg CO2e)'].sum()),
            "avg_distance": float(detailed_results['距離(公里)'].mean()),
            "top_emitters": []
        }
        
        # 按單位統計
        for unit, group in detailed_results.groupby('單位名稱'):
            analysis['units'][str(unit)] = {
                "employee_count": len(group),
                "total_emissions": float(group['總排放量(kg CO2e)'].sum()),
                "avg_distance": float(group['距離(公里)'].mean()),
                "dominant_transport": group['交通方式'].mode()[0] if len(group) > 0 else "未知"
            }
        
        # 按交通方式統計
        for transport, group in detailed_results.groupby('交通方式'):
            analysis['transport_modes'][str(transport)] = {
                "user_count": len(group),
                "total_emissions": float(group['總排放量(kg CO2e)'].sum()),
                "avg_emissions": float(group['總排放量(kg CO2e)'].mean())
            }
        
        # 找出高排放員工
        top_5 = detailed_results.nlargest(5, '總排放量(kg CO2e)')
        for _, row in top_5.iterrows():
            analysis['top_emitters'].append({
                "employee": str(row['員工名稱']),
                "unit": str(row['單位名稱']),
                "transport": str(row['交通方式']),
                "distance": float(row['距離(公里)']),
                "emissions": float(row['總排放量(kg CO2e)'])
            })
        
        self.company_profile = analysis
        self._save_json(self.company_profile_file, self.company_profile)
        
        return analysis
    
    def get_data_context(self):
        """取得數據上下文"""
        if not self.company_profile:
            return "目前沒有數據"
        
        context = f"""
## 公司碳排放數據概覽

【基本資訊】
- 員工總數: {self.company_profile.get('total_employees', 0)} 人
- 總碳排放: {self.company_profile.get('total_emissions', 0):.2f} kg CO2e
- 平均通勤距離: {self.company_profile.get('avg_distance', 0):.2f} 公里

【各單位統計】
"""
        for unit, data in self.company_profile.get('units', {}).items():
            context += f"""
{unit}:
  - 人數: {data['employee_count']} 人
  - 碳排放: {data['total_emissions']:.2f} kg CO2e
  - 平均距離: {data['avg_distance']:.2f} km
  - 主要交通方式: {data['dominant_transport']}
"""
        
        return context
    
    def save_conversation(self, user_message, ai_response):
        """儲存對話記錄"""
        conversation = {
            "timestamp": datetime.now().isoformat(),
            "user": user_message,
            "ai": ai_response
        }
        
        self.conversations.append(conversation)
        self.usage_stats["total_conversations"] += 1
        
        if len(self.conversations) > 100:
            self.conversations = self.conversations[-100:]
        
        self._save_json(self.conversation_file, self.conversations)
        self._save_json(self.usage_stats_file, self.usage_stats)
    
    def get_conversation_history(self, last_n=5):
        """取得最近的對話歷史"""
        if not self.conversations:
            return ""
        
        history = "\n## 最近對話記錄\n"
        for conv in self.conversations[-last_n:]:
            timestamp = conv['timestamp'][:19]
            history += f"\n[{timestamp}]\n"
            history += f"用戶: {conv['user']}\n"
            history += f"AI: {conv['ai'][:100]}...\n"
        
        return history
    
    def get_learning_summary(self):
        """取得學習摘要"""
        return f"""
## 系統學習狀態

【使用統計】
- 總執行次數: {self.usage_stats['total_runs']}
- 總對話次數: {self.usage_stats['total_conversations']}
- 首次使用: {self.usage_stats.get('first_use', '未知')[:10]}
- 最後使用: {self.usage_stats.get('last_use', '未知')[:10]}

【知識庫】
- 對話歷史: {len(self.conversations)} 條
- 問答資料庫: {len(self.qa_database)} 組

【公司檔案】
- 資料狀態: {'已建立' if self.company_profile else '未建立'}
- 最後更新: {self.company_profile.get('last_updated', '從未')[:19] if self.company_profile else '從未'}
"""

# ========【工具函數】========

def is_na_value(value):
    """檢查是否為NA值"""
    if pd.isna(value):
        return True
    if value is None:
        return True
    value_str = str(value).strip().upper()
    if value_str == '' or value_str == 'NA' or value_str == 'NAN':
        return True
    return False

@st.cache_resource
def get_geolocator():
    """取得地理編碼器（快取）"""
    return Nominatim(user_agent="carbon_emission_web_app")

def get_coordinates(address):
    """取得地點座標"""
    if is_na_value(address):
        return None
    try:
        geolocator = get_geolocator()
        location = geolocator.geocode(str(address), timeout=10)
        return (location.latitude, location.longitude) if location else None
    except:
        return None

def auto_fix_excel_data(df):
    """自動修復Excel資料"""
    df_fixed = df.copy()
    fix_log = []
    error_log = []
    
    transport_keywords = ['汽車', '火車', '捷運', '高鐵', '公車', '大客車', 
                          '摩托車', '機車', '飛機', '船', '電動']
    
    # 檢查欄位錯置
    swap_count = 0
    for idx, row in df_fixed.iterrows():
        unit_value = str(row['單位名稱'])
        transport_value = str(row['交通方式'])
        
        unit_has_transport = any(keyword in unit_value for keyword in transport_keywords)
        transport_is_code = len(transport_value) <= 3
        
        if unit_has_transport and transport_is_code:
            df_fixed.at[idx, '單位名稱'] = transport_value
            df_fixed.at[idx, '交通方式'] = unit_value
            fix_log.append(f"✓ 員工 {row['員工名稱']}: 已交換欄位")
            swap_count += 1
    
    # 檢查出發點NA
    for idx, row in df_fixed.iterrows():
        if is_na_value(row['出發點']):
            error_log.append(f"員工 {row['員工名稱']}: ❌ 出發點為NA")
    
    # 填補目的地NA
    dest_na_indices = []
    for idx, row in df_fixed.iterrows():
        if is_na_value(row['目的地']):
            dest_na_indices.append((idx, row['員工名稱']))
    
    if dest_na_indices:
        non_na_dests = [row['目的地'] for _, row in df_fixed.iterrows() if not is_na_value(row['目的地'])]
        if non_na_dests:
            unique_dests = list(set(non_na_dests))
            if len(unique_dests) == 1:
                fill_value = unique_dests[0]
                for idx, name in dest_na_indices:
                    df_fixed.at[idx, '目的地'] = fill_value
                fix_log.append(f"✓ 目的地: 已填補 {len(dest_na_indices)} 個NA")
    
    return df_fixed, fix_log, error_log

def calculate_emissions(distance, fuel_data, fuel_type, transport_mode):
    """計算碳排放"""
    if transport_mode in FIXED_CO2_FACTORS:
        co2 = distance * FIXED_CO2_FACTORS[transport_mode]
        details = f"距離({distance:.2f}km) × CO₂係數({FIXED_CO2_FACTORS[transport_mode]}) = {co2:.2f} kg CO₂"
        return {"CO2": co2, "CH4": 0, "N2O": 0, "Total": co2, "Details": details}
    
    matched_fuel = fuel_data[fuel_data["燃料別"] == fuel_type]
    if not matched_fuel.empty:
        co2_factor = matched_fuel.iloc[0]["CO2"]
        ch4_factor = matched_fuel.iloc[0]["CH4"]
        n2o_factor = matched_fuel.iloc[0]["N2O"]
        
        co2 = distance * co2_factor
        ch4 = distance * ch4_factor * GWP_CH4
        n2o = distance * n2o_factor * GWP_N2O
        total = co2 + ch4 + n2o
        
        details = (
            f"CO2: {distance:.2f}km × {co2_factor} = {co2:.2f} kg\n"
            f"CH4: {distance:.2f}km × {ch4_factor} × {GWP_CH4} = {ch4:.2f} kg\n"
            f"N2O: {distance:.2f}km × {n2o_factor} × {GWP_N2O} = {n2o:.2f} kg\n"
            f"總計 = {total:.2f} kg CO₂e"
        )
        return {"CO2": co2, "CH4": ch4, "N2O": n2o, "Total": total, "Details": details}
    
    return {"CO2": 0, "CH4": 0, "N2O": 0, "Total": 0, "Details": "無資料"}

def chat_with_ai(message, api_key, learning_system):
    """AI對話"""
    if not api_key:
        return "請先在側邊欄設定 OpenAI API 金鑰"
    
    try:
        # 碳排放專家提示詞
        base_prompt = """你是碳排放計算和減碳策略的專業顧問。

精通 ISO 14064-1:2018 和 IPCC 指南。

請提供具體、可行、量化的建議。"""
        
        # 添加數據上下文
        data_context = learning_system.get_data_context()
        if data_context != "目前沒有數據":
            base_prompt += f"\n\n{data_context}"
        
        # 添加對話歷史
        conversation_history = learning_system.get_conversation_history(last_n=3)
        if conversation_history:
            base_prompt += f"\n\n{conversation_history}"
        
        # 呼叫 OpenAI API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        data = {
            "model": "gpt-3.5-turbo",
            "messages": [
                {"role": "system", "content": base_prompt},
                {"role": "user", "content": message}
            ],
            "max_tokens": 800,
            "temperature": 0.3
        }
        
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content']
            
            # 儲存對話
            learning_system.save_conversation(message, ai_response)
            
            return ai_response
        else:
            return f"API 錯誤：{response.status_code}"
            
    except Exception as e:
        return f"錯誤：{str(e)}"

def generate_word_report(detailed_results, report_type='simple'):
    """生成Word報表"""
    doc = Document()
    doc.add_heading(f"碳排放成績單（{report_type}版）", level=1)
    
    # 摘要
    total_emissions = detailed_results['總排放量(kg CO2e)'].sum()
    avg_distance = detailed_results['距離(公里)'].mean()
    
    doc.add_paragraph(f"總排放量: {total_emissions:.2f} kg CO2e")
    doc.add_paragraph(f"平均距離: {avg_distance:.2f} 公里")
    doc.add_paragraph(f"總人數: {len(detailed_results)} 人")
    doc.add_paragraph("")
    
    # 表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    headers = ["員工名稱", "單位名稱", "交通方式", "燃料種類", "距離(公里)", "總排放量(kg CO2e)"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    for _, row in detailed_results.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["員工名稱"])
        cells[1].text = str(row["單位名稱"])
        cells[2].text = str(row["交通方式"])
        cells[3].text = str(row["燃料種類"])
        cells[4].text = f"{row['距離(公里)']:.2f}"
        cells[5].text = f"{row['總排放量(kg CO2e)']:.2f}"
    
    # 儲存到BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ========【初始化】========

# 初始化 session state
if 'learning_system' not in st.session_state:
    st.session_state.learning_system = ContinuousLearningSystem()

if 'detailed_results' not in st.session_state:
    st.session_state.detailed_results = pd.DataFrame()

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# ========【主介面】========

# 標題
st.markdown('<h1 style="text-align: center; color: #2E7D32;">🌱 碳排放計算工具 v3.2</h1>', unsafe_allow_html=True)
st.markdown('<p style="text-align: center; color: #666;">智能學習版 - Web Edition</p>', unsafe_allow_html=True)

# 側邊欄
with st.sidebar:
    st.header("⚙️ 設定")
    
    # API 金鑰
    st.subheader("AI 設定")
    api_key = st.text_input(
        "OpenAI API 金鑰（選填）",
        type="password",
        help="用於 AI 對話功能"
    )
    
    if api_key:
        st.success("✅ API 金鑰已設定")
    else:
        st.info("💡 不設定也可以進行計算")
    
    # 學習狀態
    st.markdown("---")
    st.subheader("📊 使用統計")
    st.metric("計算次數", st.session_state.learning_system.usage_stats['total_runs'])
    st.metric("對話次數", st.session_state.learning_system.usage_stats['total_conversations'])
    
    # 學習狀態詳情
    if st.button("查看詳細學習狀態"):
        st.text(st.session_state.learning_system.get_learning_summary())

# 主要分頁
tab1, tab2, tab3, tab4 = st.tabs(["📊 計算", "💬 AI 對話", "📈 視覺化", "📄 報表"])

# ========【Tab 1: 計算】========
with tab1:
    st.header("📊 碳排放計算")
    
    # 檔案上傳
    uploaded_file = st.file_uploader(
        "上傳 Excel 檔案",
        type=['xlsx', 'xls'],
        help="請確保檔案包含必要欄位"
    )
    
    if uploaded_file:
        st.success(f"✅ 檔案已上傳: {uploaded_file.name}")
        
        # 顯示原始資料預覽
        with st.expander("📋 原始資料預覽"):
            try:
                preview_df = pd.read_excel(uploaded_file, sheet_name='工作表2')
                st.dataframe(preview_df.head())
            except Exception as e:
                st.error(f"無法預覽: {e}")
        
        # 計算按鈕
        if st.button("🚀 開始計算", type="primary"):
            with st.spinner("計算中..."):
                try:
                    # 讀取資料
                    transport_data = pd.read_excel(uploaded_file, sheet_name='工作表2')
                    fuel_data = pd.read_excel(uploaded_file, sheet_name='工作表6')
                    
                    # 修復資料
                    transport_data, fix_log, error_log = auto_fix_excel_data(transport_data)
                    
                    # 計算排放量
                    detailed_results_list = []
                    progress_bar = st.progress(0)
                    
                    for idx, row in transport_data.iterrows():
                        progress_bar.progress((idx + 1) / len(transport_data))
                        
                        if is_na_value(row['出發點']) or is_na_value(row['目的地']):
                            detailed_results_list.append({
                                '員工名稱': row['員工名稱'],
                                '單位名稱': row['單位名稱'],
                                '交通方式': row['交通方式'],
                                '燃料種類': '-',
                                '距離(公里)': 0,
                                '總排放量(kg CO2e)': 0,
                                '計算過程': '❌ 無法計算',
                                '錯誤訊息': '請確認出發點/目的地'
                            })
                            continue
                        
                        # 取得座標並計算距離
                        origin_coords = get_coordinates(row['出發點'])
                        destination_coords = get_coordinates(row['目的地'])
                        
                        if origin_coords and destination_coords:
                            distance = geodesic(origin_coords, destination_coords).km
                            fuel_type = RECOMMENDED_FUEL_TYPES.get(row['交通方式'], "車用汽油")
                            emissions = calculate_emissions(distance, fuel_data, fuel_type, row['交通方式'])
                            
                            detailed_results_list.append({
                                '員工名稱': row['員工名稱'],
                                '單位名稱': row['單位名稱'],
                                '交通方式': row['交通方式'],
                                '燃料種類': fuel_type,
                                '距離(公里)': round(distance, 2),
                                '總排放量(kg CO2e)': round(emissions['Total'], 2),
                                '計算過程': emissions['Details'],
                                '錯誤訊息': ''
                            })
                    
                    progress_bar.empty()
                    
                    # 儲存結果
                    st.session_state.detailed_results = pd.DataFrame(detailed_results_list)
                    
                    # 分析數據
                    st.session_state.learning_system.analyze_data(st.session_state.detailed_results)
                    
                    st.success("✅ 計算完成！")
                    st.balloons()
                    
                    # 顯示修復記錄
                    if fix_log:
                        with st.expander("🔧 資料修復記錄"):
                            for log in fix_log:
                                st.text(log)
                    
                    if error_log:
                        with st.expander("⚠️ 需要確認的項目"):
                            for log in error_log:
                                st.warning(log)
                    
                except Exception as e:
                    st.error(f"❌ 計算失敗: {e}")
        
        # 顯示結果
        if not st.session_state.detailed_results.empty:
            st.markdown("---")
            st.subheader("📊 計算結果")
            
            # 統計摘要
            col1, col2, col3 = st.columns(3)
            
            results = st.session_state.detailed_results
            
            with col1:
                total_emissions = results['總排放量(kg CO2e)'].sum()
                st.metric("總排放量", f"{total_emissions:.2f} kg CO2e")
            
            with col2:
                avg_distance = results['距離(公里)'].mean()
                st.metric("平均距離", f"{avg_distance:.2f} km")
            
            with col3:
                st.metric("總人數", f"{len(results)} 人")
            
            # 詳細結果表格
            st.markdown("### 詳細結果")
            st.dataframe(
                results[['員工名稱', '單位名稱', '交通方式', '距離(公里)', '總排放量(kg CO2e)']],
                use_container_width=True
            )
            
            # 下載CSV
            csv = results.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                "📥 下載結果 (CSV)",
                csv,
                "碳排放計算結果.csv",
                "text/csv"
            )

# ========【Tab 2: AI 對話】========
with tab2:
    st.header("💬 與 AI 互動")
    
    if not api_key:
        st.warning("⚠️ 請先在側邊欄設定 OpenAI API 金鑰")
    else:
        # 顯示數據概覽
        data_summary = st.session_state.learning_system.get_data_context()
        if data_summary != "目前沒有數據":
            with st.expander("📊 AI已載入的數據摘要"):
                st.text(data_summary[:500] + "...")
        
        # 對話記錄
        for chat in st.session_state.chat_history:
            with st.chat_message(chat['role']):
                st.write(chat['content'])
        
        # 輸入框
        user_input = st.chat_input("輸入您的問題...")
        
        if user_input:
            # 顯示用戶訊息
            with st.chat_message("user"):
                st.write(user_input)
            
            st.session_state.chat_history.append({
                'role': 'user',
                'content': user_input
            })
            
            # 獲取 AI 回應
            with st.chat_message("assistant"):
                with st.spinner("思考中..."):
                    response = chat_with_ai(user_input, api_key, st.session_state.learning_system)
                    st.write(response)
            
            st.session_state.chat_history.append({
                'role': 'assistant',
                'content': response
            })
            
            st.rerun()

# ========【Tab 3: 視覺化】========
with tab3:
    st.header("📈 數據視覺化")
    
    if not st.session_state.detailed_results.empty:
        results = st.session_state.detailed_results
        
        # 圓餅圖：各單位排放占比
        st.subheader("🥧 各單位碳排放占比")
        
        unit_emissions = results.groupby('單位名稱')['總排放量(kg CO2e)'].sum()
        
        fig, ax = plt.subplots(figsize=(10, 6))
        colors = plt.cm.Set3(range(len(unit_emissions)))
        
        wedges, texts, autotexts = ax.pie(
            unit_emissions.values,
            labels=unit_emissions.index,
            autopct='%1.1f%%',
            colors=colors,
            startangle=90
        )
        
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_weight('bold')
        
        ax.set_title('各單位碳排放占比', fontsize=14, weight='bold')
        plt.tight_layout()
        st.pyplot(fig)
        
        # 長條圖：各交通方式排放量
        st.subheader("📊 各交通方式碳排放量")
        
        transport_emissions = results.groupby('交通方式')['總排放量(kg CO2e)'].sum().sort_values(ascending=False)
        
        fig2, ax2 = plt.subplots(figsize=(10, 6))
        bars = ax2.bar(transport_emissions.index, transport_emissions.values, color='#2E7D32')
        
        ax2.set_xlabel('交通方式', fontsize=12)
        ax2.set_ylabel('總排放量 (kg CO2e)', fontsize=12)
        ax2.set_title('各交通方式碳排放量', fontsize=14, weight='bold')
        
        for bar in bars:
            height = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2., height,
                    f'{height:.1f}',
                    ha='center', va='bottom')
        
        plt.xticks(rotation=45)
        plt.tight_layout()
        st.pyplot(fig2)
        
    else:
        st.info("📊 請先在「計算」頁面完成計算")

# ========【Tab 4: 報表】========
with tab4:
    st.header("📄 報表生成")
    
    if not st.session_state.detailed_results.empty:
        # 報表類型選擇
        report_type = st.radio(
            "選擇報表類型",
            ["簡易版", "詳細版"],
            horizontal=True
        )
        
        if st.button("生成 Word 報表"):
            with st.spinner("生成中..."):
                try:
                    doc_bio = generate_word_report(
                        st.session_state.detailed_results,
                        report_type
                    )
                    
                    st.download_button(
                        "📥 下載 Word 報表",
                        doc_bio,
                        f"碳排放報表_{report_type}_{datetime.now().strftime('%Y%m%d')}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ 報表已生成！")
                    
                except Exception as e:
                    st.error(f"❌ 生成失敗: {e}")
    else:
        st.info("📊 請先在「計算」頁面完成計算")

# 頁腳
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    <p>碳排放計算工具 v3.2 - Web版本</p>
    <p>© 2024 楊勢賢 | 中原大學工業與系統工程學系</p>
</div>
""", unsafe_allow_html=True)
